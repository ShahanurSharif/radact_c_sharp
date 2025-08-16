"""
Test and Demo Script for Azure AI Document Redaction
Run this script to test the functionality with sample data
"""

import sys
import os
from pathlib import Path
import structlog

# Setup logging
structlog.configure(
    processors=[
        structlog.stdlib.filter_by_level,
        structlog.stdlib.add_logger_name,
        structlog.stdlib.add_log_level,
        structlog.stdlib.PositionalArgumentsFormatter(),
        structlog.processors.TimeStamper(fmt="iso"),
        structlog.dev.ConsoleRenderer()
    ],
    wrapper_class=structlog.stdlib.BoundLogger,
    logger_factory=structlog.stdlib.LoggerFactory(),
    cache_logger_on_first_use=True,
)

logger = structlog.get_logger(__name__)

def test_azure_imports():
    """Test if Azure AI modules can be imported"""
    print("🧪 Testing Azure AI imports...")
    try:
        import azure.ai.textanalytics
        import azure.identity
        import azure.core
        print("✅ Azure AI modules imported successfully")
        return True
    except ImportError as e:
        print(f"⚠️  Azure AI modules not available: {e}")
        return False

def test_demo_redactor():
    """Test the demo redactor with sample text"""
    print("\n🔧 Testing Demo Redactor...")
    
    from demo_redactor import DemoAzureAIRedactor
    
    # Sample text with PII
    sample_text = """
    Confidential Internal Memo
    From: Jonathan Miller - Senior Security Analyst
    To: Operations Team
    
    Following our security audit, we found the following issues:
    
    1. Customer data exposure:
    - Sarah Jenkins (CEO, J-Tech Solutions) - sarah.jenkins@jtech.com
    - Michael Zhang (CFO, BrightView Capital) - m.zhang@brightview.com
    - Contact: 555-123-4567
    
    2. System vulnerabilities:
    - Server at 192.168.1.100 has exposed admin interface
    - Credit card 4532-1234-5678-9012 found in logs
    - Address: 123 Main Street, New York compromised
    
    Immediate action required.
    """
    
    redactor = DemoAzureAIRedactor()
    
    # Detect entities
    print("🔍 Detecting PII entities...")
    entities = redactor.detect_pii_entities(sample_text)
    
    print(f"Found {len(entities)} PII entities:")
    for entity in entities:
        print(f"  - {entity.category}: '{entity.text}' (confidence: {entity.confidence_score:.2f})")
    
    # Perform redaction
    print("\n🔒 Performing redaction...")
    result = redactor.redact_text(sample_text)
    
    print("📄 Original text length:", len(result.original_text))
    print("📄 Redacted text length:", len(result.redacted_text))
    print("🔢 Redactions made:", result.redaction_count)
    
    print("\n--- Redacted Text ---")
    print(result.redacted_text)
    
    # Risk analysis
    print("\n📊 Risk Analysis...")
    risk = redactor.analyze_document_risk(sample_text)
    
    print(f"Risk Score: {risk['risk_score']}")
    print(f"Risk Level: {risk['risk_level']}")
    print(f"Total Entities: {risk['total_entities']}")
    
    if risk['category_breakdown']:
        print("Category Breakdown:")
        for category, count in risk['category_breakdown'].items():
            print(f"  - {category}: {count}")
    
    if risk['recommendations']:
        print("Recommendations:")
        for rec in risk['recommendations']:
            print(f"  • {rec}")
    
    return True

def test_document_processing():
    """Test document processing with existing sample files"""
    print("\n📄 Testing Document Processing...")
    
    # Look for sample documents in the docs directory
    docs_dir = Path(__file__).parent.parent.parent / "docs"
    
    if not docs_dir.exists():
        print("⚠️  Sample documents directory not found")
        return False
    
    # Find sample documents
    docx_files = list(docs_dir.glob("*.docx"))
    pdf_files = list(docs_dir.glob("*.pdf"))
    
    print(f"Found {len(docx_files)} DOCX files and {len(pdf_files)} PDF files")
    
    if docx_files or pdf_files:
        print("📁 Sample documents available for processing:")
        for doc in docx_files + pdf_files:
            print(f"  - {doc.name}")
        
        print("\n💡 To process these documents, run:")
        print(f"python main.py analyze {docs_dir / '1.docx'}")
        print("(Note: This requires Azure credentials)")
    
    return True

def create_sample_documents():
    """Create sample documents for testing"""
    print("\n📝 Creating sample documents...")
    
    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        
        # Sample text with PII
        sample_content = """
        CONFIDENTIAL EMPLOYEE RECORD
        
        Employee: Sarah Johnson
        Position: Senior Manager
        Email: sarah.johnson@company.com
        Phone: (555) 123-4567
        Address: 456 Oak Street, Boston, MA 02101
        
        Emergency Contact: Michael Johnson (spouse)
        Emergency Phone: (555) 987-6543
        
        Credit Card (company): 4532-1234-5678-9012
        SSN: 123-45-6789
        
        This document contains sensitive information and should be handled according to privacy policies.
        """
        
        # Create DOCX
        doc = Document()
        doc.add_heading('Sample Document for Redaction Testing', 0)
        doc.add_paragraph(sample_content)
        
        docx_path = Path('sample_document.docx')
        doc.save(docx_path)
        print(f"✅ Created sample DOCX: {docx_path}")
        
        # Create PDF
        pdf_path = Path('sample_document.pdf')
        pdf_doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
        styles = getSampleStyleSheet()
        
        story = []
        story.append(Paragraph('Sample Document for Redaction Testing', styles['Title']))
        
        for paragraph in sample_content.split('\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph, styles['Normal']))
        
        pdf_doc.build(story)
        print(f"✅ Created sample PDF: {pdf_path}")
        
        return True
        
    except ImportError as e:
        print(f"⚠️  Cannot create sample documents: {e}")
        return False

def run_cli_demo():
    """Demonstrate CLI functionality"""
    print("\n🖥️  CLI Demo Commands:")
    print("=" * 40)
    
    print("1. Create configuration file:")
    print("   python main.py config")
    
    print("\n2. Analyze a document:")
    print("   python main.py analyze sample_document.docx")
    
    print("\n3. Redact a single document:")
    print("   python main.py file sample_document.docx")
    
    print("\n4. Batch process a directory:")
    print("   python main.py directory ./documents")
    
    print("\n5. Get help:")
    print("   python main.py --help")

def main():
    """Main demo function"""
    print("🚀 Azure AI Document Redaction - Test & Demo")
    print("=" * 50)
    
    # Test imports
    azure_available = test_azure_imports()
    
    # Test demo functionality
    demo_success = test_demo_redactor()
    
    # Test document processing capabilities
    doc_test = test_document_processing()
    
    # Create sample documents
    sample_created = create_sample_documents()
    
    # Show CLI demo
    run_cli_demo()
    
    print("\n📋 Summary:")
    print(f"✅ Azure AI Available: {'Yes' if azure_available else 'No (Demo mode only)'}")
    print(f"✅ Demo Functionality: {'Working' if demo_success else 'Failed'}")
    print(f"✅ Document Processing: {'Ready' if doc_test else 'Limited'}")
    print(f"✅ Sample Documents: {'Created' if sample_created else 'Not created'}")
    
    if not azure_available:
        print("\n🔧 To enable Azure AI features:")
        print("1. Set up Azure Text Analytics service")
        print("2. Edit .env file with your credentials")
        print("3. Run: python main.py analyze sample_document.docx")
    
    print("\n🎉 Demo completed! Ready for document redaction.")

if __name__ == "__main__":
    main()
