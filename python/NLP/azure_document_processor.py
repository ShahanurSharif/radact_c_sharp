"""
Document processors with Azure AI integration
Handles DOCX and PDF files with intelligent redaction
"""

import os
import sys
from typing import Optional, Dict, Any
from pathlib import Path

# Add parent directory to path for imports
sys.path.append(str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import PyPDF2
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import structlog

from azure_ai_redactor import AzureAIRedactor, RedactionResult
from azure_config import AzureConfig

logger = structlog.get_logger(__name__)

class AzureDocxProcessor:
    """DOCX processor with Azure AI redaction capabilities"""
    
    def __init__(self, config: Optional[AzureConfig] = None):
        """
        Initialize DOCX processor with Azure AI
        
        Args:
            config: Azure configuration
        """
        self.redactor = AzureAIRedactor(config)
        logger.info("Azure DOCX processor initialized")
    
    def process_document(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """
        Process DOCX document with Azure AI redaction
        
        Args:
            input_path: Path to input DOCX file
            output_path: Path to output redacted DOCX file
            
        Returns:
            Processing results and statistics
        """
        try:
            # Load document
            doc = Document(input_path)
            
            # Extract all text content
            full_text = ""
            paragraph_texts = []
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text
                paragraph_texts.append(para_text)
                full_text += para_text + "\n"
            
            logger.info("Document loaded", 
                       paragraphs=len(paragraph_texts),
                       total_chars=len(full_text))
            
            # Perform Azure AI redaction on full text
            redaction_result = self.redactor.redact_text(full_text)
            
            # Split redacted text back into paragraphs
            redacted_paragraphs = redaction_result.redacted_text.split('\n')
            
            # Apply redactions to document
            for i, paragraph in enumerate(doc.paragraphs):
                if i < len(redacted_paragraphs):
                    # Clear existing text
                    paragraph.clear()
                    
                    # Add redacted text with formatting
                    run = paragraph.add_run(redacted_paragraphs[i])
                    
                    # Highlight redacted content
                    if '[' in redacted_paragraphs[i] and '_REDACTED]' in redacted_paragraphs[i]:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red text
            
            # Add redaction summary at the beginning
            summary_paragraph = doc.paragraphs[0]
            summary_paragraph.insert_paragraph_before(
                f"🔒 DOCUMENT REDACTED - {redaction_result.redaction_count} entities redacted"
            ).runs[0].font.bold = True
            
            # Save redacted document
            doc.save(output_path)
            
            # Analyze document risk
            risk_analysis = self.redactor.analyze_document_risk(full_text)
            
            results = {
                'status': 'success',
                'input_file': input_path,
                'output_file': output_path,
                'original_length': len(full_text),
                'redacted_length': len(redaction_result.redacted_text),
                'entities_found': len(redaction_result.entities_found),
                'redactions_made': redaction_result.redaction_count,
                'confidence_scores': redaction_result.confidence_scores,
                'risk_analysis': risk_analysis,
                'entity_details': [
                    {
                        'text': entity.text,
                        'category': entity.category,
                        'confidence': entity.confidence_score
                    }
                    for entity in redaction_result.entities_found
                ]
            }
            
            logger.info("DOCX processing completed", results=results)
            return results
            
        except Exception as e:
            error_msg = f"Error processing DOCX: {str(e)}"
            logger.error("DOCX processing failed", error=error_msg)
            return {
                'status': 'error',
                'error': error_msg,
                'input_file': input_path
            }

class AzurePdfProcessor:
    """PDF processor with Azure AI redaction capabilities"""
    
    def __init__(self, config: Optional[AzureConfig] = None):
        """
        Initialize PDF processor with Azure AI
        
        Args:
            config: Azure configuration
        """
        self.redactor = AzureAIRedactor(config)
        logger.info("Azure PDF processor initialized")
    
    def process_document(self, input_path: str, output_path: str) -> Dict[str, Any]:
        """
        Process PDF document with Azure AI redaction
        
        Args:
            input_path: Path to input PDF file
            output_path: Path to output redacted PDF file
            
        Returns:
            Processing results and statistics
        """
        try:
            # Extract text from PDF
            full_text = ""
            with open(input_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    full_text += page_text + "\n"
            
            logger.info("PDF text extracted", 
                       pages=len(pdf_reader.pages),
                       total_chars=len(full_text))
            
            # Perform Azure AI redaction
            redaction_result = self.redactor.redact_text(full_text)
            
            # Create new PDF with redacted content
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Custom style for redacted content
            redacted_style = ParagraphStyle(
                'RedactedContent',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=12,
                leftIndent=0.5*inch,
                rightIndent=0.5*inch
            )
            
            story = []
            
            # Add redaction header
            header_style = ParagraphStyle(
                'RedactionHeader',
                parent=styles['Title'],
                fontSize=14,
                spaceAfter=20,
                textColor='red'
            )
            
            story.append(Paragraph(
                f"🔒 DOCUMENT REDACTED - {redaction_result.redaction_count} entities redacted",
                header_style
            ))
            story.append(Spacer(1, 20))
            
            # Add redacted content
            paragraphs = redaction_result.redacted_text.split('\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para, redacted_style))
                    story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
            # Analyze document risk
            risk_analysis = self.redactor.analyze_document_risk(full_text)
            
            results = {
                'status': 'success',
                'input_file': input_path,
                'output_file': output_path,
                'original_length': len(full_text),
                'redacted_length': len(redaction_result.redacted_text),
                'entities_found': len(redaction_result.entities_found),
                'redactions_made': redaction_result.redaction_count,
                'confidence_scores': redaction_result.confidence_scores,
                'risk_analysis': risk_analysis,
                'entity_details': [
                    {
                        'text': entity.text,
                        'category': entity.category,
                        'confidence': entity.confidence_score
                    }
                    for entity in redaction_result.entities_found
                ]
            }
            
            logger.info("PDF processing completed", results=results)
            return results
            
        except Exception as e:
            error_msg = f"Error processing PDF: {str(e)}"
            logger.error("PDF processing failed", error=error_msg)
            return {
                'status': 'error',
                'error': error_msg,
                'input_file': input_path
            }

class AzureDocumentProcessor:
    """Main document processor that handles both DOCX and PDF files"""
    
    def __init__(self, config: Optional[AzureConfig] = None):
        """
        Initialize document processor with Azure AI
        
        Args:
            config: Azure configuration
        """
        self.config = config or AzureConfig()
        self.docx_processor = AzureDocxProcessor(self.config)
        self.pdf_processor = AzurePdfProcessor(self.config)
        
        logger.info("Azure document processor initialized")
    
    def process_file(self, input_path: str, output_path: Optional[str] = None) -> Dict[str, Any]:
        """
        Process document file with Azure AI redaction
        
        Args:
            input_path: Path to input document
            output_path: Path to output document (auto-generated if None)
            
        Returns:
            Processing results
        """
        input_file = Path(input_path)
        
        if not input_file.exists():
            return {
                'status': 'error',
                'error': f'Input file not found: {input_path}'
            }
        
        # Auto-generate output path if not provided
        if output_path is None:
            output_path = str(input_file.parent / f"azure_redacted_{input_file.name}")
        
        # Determine file type and process accordingly
        file_extension = input_file.suffix.lower()
        
        if file_extension == '.docx':
            return self.docx_processor.process_document(input_path, output_path)
        elif file_extension == '.pdf':
            return self.pdf_processor.process_document(input_path, output_path)
        else:
            return {
                'status': 'error',
                'error': f'Unsupported file type: {file_extension}',
                'supported_types': ['.docx', '.pdf']
            }
    
    def batch_process(self, input_directory: str, output_directory: Optional[str] = None) -> Dict[str, Any]:
        """
        Process multiple documents in a directory
        
        Args:
            input_directory: Directory containing input documents
            output_directory: Directory for output documents
            
        Returns:
            Batch processing results
        """
        input_dir = Path(input_directory)
        
        if not input_dir.exists():
            return {
                'status': 'error',
                'error': f'Input directory not found: {input_directory}'
            }
        
        if output_directory is None:
            output_directory = str(input_dir / "azure_redacted")
        
        output_dir = Path(output_directory)
        output_dir.mkdir(exist_ok=True)
        
        # Find supported documents
        supported_extensions = ['.docx', '.pdf']
        documents = []
        
        for ext in supported_extensions:
            documents.extend(input_dir.glob(f'*{ext}'))
        
        results = []
        success_count = 0
        error_count = 0
        
        for doc_path in documents:
            output_path = str(output_dir / f"azure_redacted_{doc_path.name}")
            result = self.process_file(str(doc_path), output_path)
            
            results.append(result)
            
            if result['status'] == 'success':
                success_count += 1
            else:
                error_count += 1
        
        batch_results = {
            'status': 'completed',
            'input_directory': input_directory,
            'output_directory': output_directory,
            'documents_processed': len(documents),
            'successful': success_count,
            'failed': error_count,
            'results': results
        }
        
        logger.info("Batch processing completed", summary=batch_results)
        return batch_results
