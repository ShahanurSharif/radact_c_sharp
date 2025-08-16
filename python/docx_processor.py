"""
DOCX document processor for Python version
"""

from docx import Document
from pathlib import Path
import os
from main import CommonProcessor


class DocxProcessor(CommonProcessor):
    """DOCX document processor"""
    
    def __init__(self, redactor):
        super().__init__(redactor)
    
    def read_content(self, file_path: str) -> str:
        """Read content from DOCX file"""
        try:
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
            
            return '\n'.join(content)
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {e}")
    
    def generate_output_path(self, input_file_path: str) -> str:
        """Generate output path for redacted DOCX file"""
        file_path = Path(input_file_path)
        file_name = file_path.stem
        output_dir = Path("../docs")
        return str(output_dir / f"redact_{file_name}.docx")
    
    def save_content(self, input_file_path: str, output_path: str, redacted_content: str):
        """Save redacted content as DOCX file"""
        try:
            # Create new document
            doc = Document()
            
            # Add title
            title = doc.add_paragraph("REDACTED DOCUMENT")
            title.alignment = 1  # Center alignment
            
            # Add redacted content
            lines = redacted_content.split('\n')
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph("")  # Empty line
            
            # Save the document
            doc.save(output_path)
            print(f"✅ DOCX successfully created: {output_path}")
            
        except Exception as e:
            print(f"❌ Error creating DOCX: {e}")
            # Fallback: save as text
            text_path = output_path.replace(".docx", "_fallback.txt")
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(redacted_content)
            print(f"💾 Saved as text file instead: {text_path}")
