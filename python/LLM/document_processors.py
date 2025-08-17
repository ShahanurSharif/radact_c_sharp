"""
Document processors for LLM-based redaction
Handles DOCX and PDF documents with GPT-4o-mini integration
"""

import os
from pathlib import Path
from typing import Optional, Dict, Any
from dataclasses import dataclass
import structlog

# Document processing imports
from docx import Document
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
import io

from gpt4o_redactor import GPT4oMiniRedactor, RedactionResult
from llm_config import LLMConfig
from enhanced_pdf_processor import EnhancedPdfProcessor

logger = structlog.get_logger(__name__)

@dataclass 
class DocumentInfo:
    """Information about processed document"""
    file_path: str
    file_type: str
    page_count: int
    word_count: int
    character_count: int
    processing_cost: float
    entities_found: int
    risk_level: str

class DocumentProcessor:
    """Base document processor with LLM redaction capabilities"""
    
    def __init__(self, config: Optional[LLMConfig] = None):
        """
        Initialize document processor
        
        Args:
            config: LLM configuration for redaction
        """
        self.config = config or LLMConfig()
        self.redactor = GPT4oMiniRedactor(self.config)
        
        logger.info("Document processor initialized with LLM redaction")
    
    def _assess_risk_level(self, entities_count: int, confidence_scores: Dict[str, float]) -> str:
        """Assess document risk level based on PII findings"""
        
        if entities_count == 0:
            return "LOW"
        
        # High-risk categories
        high_risk_categories = ['ssn', 'credit_cards']
        has_high_risk = any(cat in confidence_scores for cat in high_risk_categories)
        
        # Risk assessment logic
        if has_high_risk or entities_count >= 10:
            return "HIGH"
        elif entities_count >= 5:
            return "MEDIUM"
        else:
            return "LOW"
    
    def _create_redacted_filename(self, original_path: str, suffix: str = "_redacted") -> str:
        """Create filename for redacted document"""
        
        path = Path(original_path)
        new_name = f"{path.stem}{suffix}{path.suffix}"
        return str(path.parent / new_name)
    
    def get_cost_estimate(self, file_path: str) -> Dict[str, Any]:
        """
        Get cost estimate for processing document
        
        Args:
            file_path: Path to document
            
        Returns:
            Cost estimation details
        """
        try:
            # Extract text based on file type
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.docx':
                text = self._extract_docx_text(file_path)
                # Get LLM cost estimate
                estimate = self.redactor.get_cost_estimate(text)
                
                # Add document info
                estimate.update({
                    'file_path': file_path,
                    'file_type': file_ext,
                    'text_length': len(text),
                    'word_count': len(text.split())
                })
                
            elif file_ext == '.pdf':
                # Use enhanced PDF processor for cost estimation
                enhanced_processor = EnhancedPdfProcessor(self.config)
                estimate = enhanced_processor.get_cost_estimate(file_path)
                
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            return estimate
            
        except Exception as e:
            logger.error("Failed to estimate cost", file_path=file_path, error=str(e))
            return {'error': str(e)}
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error("Failed to extract DOCX text", file_path=file_path, error=str(e))
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        
        try:
            text_parts = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.error("Failed to extract PDF text", file_path=file_path, error=str(e))
            raise

class DocxProcessor(DocumentProcessor):
    """DOCX document processor with LLM redaction"""
    
    def process_document(self, file_path: str, output_path: Optional[str] = None) -> DocumentInfo:
        """
        Process DOCX document with LLM-based redaction
        
        Args:
            file_path: Path to input DOCX file
            output_path: Optional output path for redacted document
            
        Returns:
            Document processing information
        """
        logger.info("Processing DOCX document", file_path=file_path)
        
        if output_path is None:
            output_path = self._create_redacted_filename(file_path)
        
        try:
            # Load document
            doc = Document(file_path)
            
            # Extract text for redaction
            full_text = self._extract_docx_text(file_path)
            
            # Perform LLM redaction
            redaction_result = self.redactor.redact_text(full_text)
            
            # Apply redactions to document
            self._apply_docx_redactions(doc, redaction_result)
            
            # Save redacted document
            doc.save(output_path)
            
            # Calculate document stats
            word_count = len(full_text.split())
            char_count = len(full_text)
            page_count = self._estimate_docx_pages(doc)
            
            # Assess risk level
            risk_level = self._assess_risk_level(
                redaction_result.total_entities,
                redaction_result.confidence_scores
            )
            
            doc_info = DocumentInfo(
                file_path=output_path,
                file_type='docx',
                page_count=page_count,
                word_count=word_count,
                character_count=char_count,
                processing_cost=redaction_result.processing_cost,
                entities_found=redaction_result.total_entities,
                risk_level=risk_level
            )
            
            logger.info("DOCX processing completed",
                       output_path=output_path,
                       entities_redacted=redaction_result.total_entities,
                       cost=redaction_result.processing_cost,
                       risk_level=risk_level)
            
            return doc_info
            
        except Exception as e:
            logger.error("Failed to process DOCX", file_path=file_path, error=str(e))
            raise
    
    def _apply_docx_redactions(self, doc: Document, redaction_result: RedactionResult):
        """Apply redactions to DOCX document while preserving formatting"""
        
        # Get redaction patterns from config
        redaction_patterns = self.config.get_redaction_patterns()
        
        # Process each paragraph individually to preserve structure
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
                
            # Apply redactions to this paragraph's text
            original_text = paragraph.text
            redacted_text = original_text
            
            # Apply redactions for each entity found
            for entity in redaction_result.entities_found:
                if entity.text in original_text:
                    # Get appropriate redaction pattern
                    pattern = redaction_patterns.get(entity.category, '[REDACTED]')
                    redacted_text = redacted_text.replace(entity.text, pattern)
                    
                    logger.debug("Applied DOCX redaction",
                               original=entity.text,
                               replacement=pattern,
                               category=entity.category)
            
            # Update paragraph text while preserving formatting
            if redacted_text != original_text:
                # Clear existing runs and add redacted text
                paragraph.clear()
                paragraph.add_run(redacted_text)
        
        logger.info("Applied DOCX redactions", 
                   total_entities=redaction_result.total_entities)
    
    def _estimate_docx_pages(self, doc: Document) -> int:
        """Estimate number of pages in DOCX document"""
        
        # Rough estimation: ~500 words per page
        total_words = sum(len(p.text.split()) for p in doc.paragraphs)
        return max(1, (total_words + 499) // 500)

class PdfProcessor(DocumentProcessor):
    """PDF document processor with LLM redaction and format preservation"""
    
    def process_document(self, file_path: str, output_path: Optional[str] = None) -> DocumentInfo:
        """
        Process PDF document with LLM-based redaction while preserving formatting
        
        Args:
            file_path: Path to input PDF file
            output_path: Optional output path for redacted PDF
            
        Returns:
            Document processing information
        """
        logger.info("Processing PDF document with format preservation", file_path=file_path)
        
        try:
            # Use enhanced PDF processor that preserves formatting
            enhanced_processor = EnhancedPdfProcessor(self.config)
            result_info = enhanced_processor.process_pdf(file_path, output_path)
            
            # Convert to DocumentInfo format
            doc_info = DocumentInfo(
                file_path=result_info['file_path'],
                file_type=result_info['file_type'],
                page_count=result_info['page_count'],
                word_count=result_info['word_count'],
                character_count=result_info['character_count'],
                processing_cost=result_info['processing_cost'],
                entities_found=result_info['entities_found'],
                risk_level=result_info['risk_level']
            )
            
            logger.info("PDF processing completed with format preservation",
                       output_path=result_info['file_path'],
                       entities_redacted=result_info['entities_found'],
                       cost=result_info['processing_cost'],
                       risk_level=result_info['risk_level'])
            
            return doc_info
            
        except Exception as e:
            logger.error("Failed to process PDF", file_path=file_path, error=str(e))
            raise
    
    def _create_redacted_pdf(self, redacted_text: str, output_path: str):
        """Create new PDF with redacted content"""
        
        try:
            # Create PDF with reportlab
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            
            # Define styles
            styles = getSampleStyleSheet()
            normal_style = styles['Normal']
            
            # Split text into paragraphs
            paragraphs = redacted_text.split('\n\n')
            
            # Create PDF content
            story = []
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    # Create paragraph
                    para = Paragraph(paragraph_text, normal_style)
                    story.append(para)
                    story.append(Spacer(1, 0.1 * inch))
            
            # Build PDF
            doc.build(story)
            
        except Exception as e:
            logger.error("Failed to create redacted PDF", output_path=output_path, error=str(e))
            raise

def create_processor(file_path: str, config: Optional[LLMConfig] = None) -> DocumentProcessor:
    """
    Factory function to create appropriate document processor
    
    Args:
        file_path: Path to document
        config: LLM configuration
        
    Returns:
        Document processor instance
    """
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.docx':
        return DocxProcessor(config)
    elif file_ext == '.pdf':
        return PdfProcessor(config)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")
